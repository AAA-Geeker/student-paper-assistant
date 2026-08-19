"""从上传文件提取纯文本，供核心功能导入 Word/PDF 使用。

支持：
- .docx  → python-docx（已存在依赖）
- .pdf   → pypdf（轻量，纯 Python，无原生依赖，适合腾讯云构建）
- .txt   → UTF-8 解码
- .md    → 按纯文本读取
不支持或解析失败时抛 ValueError（由路由转成 400 友好提示）。
"""
from io import BytesIO

from fastapi import HTTPException


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF 解析组件未安装，请联系管理员")
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"PDF 解析失败：{e}")
    pages = []
    for p in reader.pages:
        try:
            pages.append(p.extract_text() or "")
        except Exception as e:  # noqa: BLE001
            pages.append(f"[第 {len(pages)+1} 页提取失败：{e}]")
    return "\n\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    from docx import Document

    try:
        doc = Document(BytesIO(content))
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"Word 文档解析失败：{e}")
    parts = []
    # 段落文字（含表格外）
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格内文字（论文常含表格）
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise HTTPException(status_code=400, detail="Word 文档中未提取到文字，请确认文档非空白或非图片型")
    return text


def extract_text_from_file(filename: str, content: bytes) -> str:
    """按文件扩展名提取纯文本。filename 可能带路径，取最后一段。"""
    name = (filename or "").rsplit("/", 1)[-1].rsplit("\\", 1)[-1].strip()
    ext = ("." in name and name.rsplit(".", 1)[-1].lower()) or ""
    if not content:
        raise HTTPException(status_code=400, detail="文件内容为空")

    if ext == "pdf":
        text = _extract_pdf(content)
    elif ext == "docx":
        text = _extract_docx(content)
    elif ext in ("txt", "md"):
        text = _decode_text(content)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"暂不支持 {ext or '无扩展名'} 格式，请上传 .docx / .pdf / .txt 文件",
        )

    if not text:
        raise HTTPException(status_code=400, detail="文件中未提取到文字，请确认内容为可选中文本（非扫描图片）")
    return text


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8").strip()
    except UnicodeDecodeError:
        pass
    for enc in ("gb18030", "gbk", "big5"):
        try:
            return content.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="文本文件编码无法识别（请另存为 UTF-8）")

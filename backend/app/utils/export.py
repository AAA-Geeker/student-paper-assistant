from io import BytesIO
from docx import Document


def to_markdown(title: str, content: str) -> str:
    return "# " + title + "\n\n" + content


def to_docx(title: str, content: str) -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in content.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# 跨平台中文字体查找：本地 Windows / macOS / Docker(Linux，fonts-noto-cjk)
_CHINESE_FONT_CANDIDATES = [
    # Windows
    r"C:\Windows\Fonts\msyh.ttc",      # 微软雅黑
    r"C:\Windows\Fonts\simhei.ttf",    # 黑体
    r"C:\Windows\Fonts\simsun.ttc",    # 宋体
    r"C:\Windows\Fonts\msyh.ttf",
    # macOS
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    # Linux (Docker: fonts-noto-cjk)
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
]

_ZH_FONT_PATH = None
_ZH_FONT_INDEX = 0  # 对 .ttc 集合字体，用第一个子字体


def _find_chinese_font() -> str:
    """在常见系统字体里找一个中文字体 TrueType/OpenType 文件路径。找不到则抛错。"""
    global _ZH_FONT_PATH
    if _ZH_FONT_PATH:
        return _ZH_FONT_PATH
    for cand in _CHINESE_FONT_CANDIDATES:
        import os
        if os.path.exists(cand):
            _ZH_FONT_PATH = cand
            if cand.lower().endswith(".ttf"):
                _ZH_FONT_INDEX = 0
            return cand
    raise RuntimeError(
        "PDF 导出需要系统中文字体，但未找到。请安装中文字体（如微软雅黑/Noto CJK），"
        "或将字体路径加入 export._CHINESE_FONT_CANDIDATES。"
    )


def to_pdf(title: str, content: str) -> BytesIO:
    """用 fpdf2 生成 PDF（纯 Python，无原生依赖），自动适配中文字体。"""
    from fpdf import FPDF

    font_path = _find_chinese_font()
    font_is_collection = font_path.lower().endswith(".ttc")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if font_is_collection:
        pdf.add_font("zh", "", font_path, uni=True)  # fpdf2 自动读取 ttc 第一个子字体
    else:
        pdf.add_font("zh", "", font_path)

    # 标题
    pdf.set_font("zh", "", 16)
    pdf.multi_cell(0, 9, title)
    pdf.ln(4)

    # 正文：按段落自然换行，保留空行
    pdf.set_font("zh", "", 12)
    for para in content.split("\n\n"):
        if para.strip():
            para = para.replace("\n", " ")
            pdf.multi_cell(0, 7, para)
        pdf.ln(2)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

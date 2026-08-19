"""文件文本提取单测：验证 core 功能导入 Word/PDF 的文本提取逻辑。

生成真实 .docx / .pdf / .txt 临时文件测试 extract_text_from_file，
不依赖真实 SMTP 或网络。此实现是对第 2 项（导入 Word/PDF）的回归保护。
"""
from io import BytesIO

import pytest
from fastapi import HTTPException

from app.utils.file_extract import extract_text_from_file


@pytest.fixture()
def docx_bytes():
    from docx import Document

    doc = Document()
    doc.add_heading("测试标题", 0)
    doc.add_paragraph("这是第一段用于降重测试的中文正文。")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "方法"
    t.rows[0].cells[1].text = "结果"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture()
def pdf_bytes():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=13)
    pdf.multi_cell(0, 8, "First PDF page review content.")
    pdf.add_page()
    pdf.multi_cell(0, 8, "Second page content.")
    buf = BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def test_extract_docx_includes_paragraphs_and_tables(docx_bytes):
    text = extract_text_from_file("paper.docx", docx_bytes)
    assert "测试标题" in text
    assert "中文正文" in text
    assert "方法" in text and "结果" in text  # 表格也需提取


def test_extract_pdf_multipage(pdf_bytes):
    text = extract_text_from_file("paper.pdf", pdf_bytes)
    assert "First PDF page" in text
    assert "Second page" in text


def test_extract_txt_utf8():
    text = extract_text_from_file("a.txt", "中文内容 123".encode("utf-8"))
    assert "中文内容" in text


def test_unsupported_extension_raises_400():
    with pytest.raises(HTTPException) as e:
        extract_text_from_file("image.png", b"\x89PNG")
    assert e.value.status_code == 400
    assert "暂不支持" in e.value.detail


def test_empty_content_raises_400():
    with pytest.raises(HTTPException) as e:
        extract_text_from_file("a.pdf", b"")
    assert e.value.status_code == 400
